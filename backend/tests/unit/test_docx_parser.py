"""
Unit tests for app.graph.nodes.parsers.docx_parser.parse_docx().

TDD red phase: written BEFORE docx_parser.py exists (Task 9).
Run: python -m pytest tests/unit/test_docx_parser.py -v
Expected before Task 10: FAIL with ImportError
Expected after  Task 10: PASS (OCR tests skipped if Tesseract absent)
"""

from unittest.mock import patch

import pytest
from docx import Document as DocxDocument

from app.graph.nodes.parsers import ParseResult

# This import will raise ImportError until Task 10 creates the module — expected.
from app.graph.nodes.parsers.docx_parser import parse_docx

# ─── Direct text extraction ────────────────────────────────────────────────────


def test_parse_docx_digital_text(sample_docx_path):
    """Standard DOCX: direct text extraction succeeds without OCR."""
    result = parse_docx(sample_docx_path, timeout_seconds=60)
    assert isinstance(result, ParseResult)
    assert len(result.text) >= 200
    assert result.ocr_used is False
    assert result.ocr_confidence is None


def test_parse_docx_page_count_heuristic(sample_docx_path):
    """Page count uses heuristic max(1, len(text) // 3000) — always >= 1."""
    result = parse_docx(sample_docx_path, timeout_seconds=60)
    # fixture has ~496 chars → max(1, 496 // 3000) = max(1, 0) = 1
    assert result.page_count >= 1


# ─── OCR fallback ──────────────────────────────────────────────────────────────


@pytest.mark.skipif(
    not __import__("shutil").which("tesseract"),
    reason="Tesseract OCR is not installed",
)
def test_parse_docx_empty_triggers_ocr(tmp_path):
    """DOCX with <50 chars triggers OCR attempt (function must not crash)."""
    doc = DocxDocument()
    doc.add_paragraph("Hi")  # 2 chars — well below 50
    path = str(tmp_path / "empty.docx")
    doc.save(path)

    # OCR attempt may or may not succeed (depends on pymupdf DOCX rendering)
    # but the function must return a ParseResult — never raise.
    result = parse_docx(path, timeout_seconds=60)
    assert isinstance(result, ParseResult)


def test_parse_docx_ocr_rendering_failure_graceful(tmp_path):
    """If pymupdf cannot render DOCX for OCR, fall back to direct text gracefully.

    Creates a DOCX whose char density is low (60 chars / 1 page = 60 < 100)
    so OCR is triggered. When fitz.open raises (e.g. unsupported variant),
    the parser must return the direct-extracted text with ocr_used=False, not raise.
    """
    doc = DocxDocument()
    doc.add_paragraph("A" * 60)  # 60 chars, 1 page → density 60 < 100
    path = str(tmp_path / "rendertest.docx")
    doc.save(path)

    # Simulate pymupdf failing to render the DOCX (the scenario named in this test)
    with patch("fitz.open", side_effect=RuntimeError("Cannot render this DOCX variant")):
        result = parse_docx(path, timeout_seconds=60)

    assert isinstance(result, ParseResult)
    assert result.ocr_used is False
    # Direct text must be preserved when OCR rendering fails
    assert "A" * 60 in result.text


# ─── Error handling ────────────────────────────────────────────────────────────


def test_parse_docx_corrupted_raises_value_error(tmp_path):
    """Corrupted DOCX (wrong binary content) raises ValueError."""
    path = str(tmp_path / "bad.docx")
    with open(path, "wb") as f:
        f.write(b"not a valid docx file at all")
    with pytest.raises(ValueError):
        parse_docx(path, timeout_seconds=60)


def test_parse_docx_timeout_raises(sample_docx_path):
    """Processing that exceeds timeout_seconds raises TimeoutError.

    Deterministic: the inner parse is patched to block for 0.3s — an order of
    magnitude longer than the 0.02s timeout — so the executor's
    future.result(timeout=...) reliably fires. The previous version relied on a
    real sub-millisecond parse losing a race against a 0.001s timeout, which was
    flaky (the tiny fixture often finished before the deadline).
    """
    import time

    def _slow_parse(_file_path):
        time.sleep(0.3)  # >> the 0.02s timeout below; cannot finish in time

    with patch(
        "app.graph.nodes.parsers.docx_parser._parse_docx_inner",
        side_effect=_slow_parse,
    ):
        with pytest.raises(TimeoutError):
            parse_docx(sample_docx_path, timeout_seconds=0.02)
